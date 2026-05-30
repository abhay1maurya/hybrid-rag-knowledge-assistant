import os
from typing import List
from langchain_core.documents import Document
from .base_handler import BaseFileHandler


class DOCXHandler(BaseFileHandler):
    supported_extensions = [".docx", ".doc"]

    def load(self, file_path: str) -> List[Document]:
        print(f"  [DOCXHandler] Loading: {file_path}")
        try:
            import docx
        except ImportError:
            raise RuntimeError("Run: pip install python-docx")

        doc       = docx.Document(file_path)
        documents = []
        filename  = os.path.basename(file_path)

        # ── Extract paragraphs grouped by heading sections ────────────────────
        sections = self._extract_sections(doc)

        for i, section in enumerate(sections):
            if section["content"].strip():
                documents.append(Document(
                    page_content=section["content"],
                    metadata={
                        "source":        file_path,
                        "file_type":     "docx",
                        "filename":      filename,
                        "section_title": section["title"],
                        "page":          i + 1,
                        "section":       i + 1,
                    }
                ))

        # ── Extract tables as separate documents ──────────────────────────────
        table_docs = self._extract_tables(doc, file_path, filename, len(documents))
        documents.extend(table_docs)

        print(f"  [DOCXHandler] Loaded {len(documents)} sections + tables.")
        return documents

    def _extract_sections(self, doc) -> list:
        """Groups paragraphs under their heading sections."""
        sections        = []
        current_title   = "Introduction"
        current_content = []

        for para in doc.paragraphs:
            text  = para.text.strip()
            style = para.style.name.lower()

            if not text:
                continue

            # New heading = new section
            if "heading" in style:
                # Save previous section
                if current_content:
                    sections.append({
                        "title":   current_title,
                        "content": "\n".join(current_content)
                    })
                current_title   = text
                current_content = []
            else:
                current_content.append(text)

        # Save last section
        if current_content:
            sections.append({
                "title":   current_title,
                "content": "\n".join(current_content)
            })

        return sections if sections else [{"title": "Content", "content": "\n".join(
            [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        )}]

    def _extract_tables(self, doc, file_path: str, filename: str, offset: int) -> List[Document]:
        """Converts tables to readable text documents."""
        table_docs = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)

            if table_text.strip():
                table_docs.append(Document(
                    page_content=f"[Table {i+1}]\n{table_text}",
                    metadata={
                        "source":    file_path,
                        "file_type": "docx",
                        "filename":  filename,
                        "page":      offset + i + 1,
                        "type":      "table",
                        "table_index": i + 1
                    }
                ))
        return table_docs